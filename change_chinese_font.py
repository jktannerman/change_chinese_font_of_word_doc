"""Change the East Asian font of Chinese characters in a .docx file.

Sets `w:eastAsia` on each run that contains CJK characters, leaving
Latin/other fonts untouched. Covers body paragraphs, tables (including
nested), headers, footers, and text boxes.

Usage:
    py -3.13 change_chinese_font.py <input.docx> [--output <out.docx>] [--font <font_name>]
"""

from __future__ import annotations

import argparse
import re
import sys
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.table import Table

DEFAULT_FONT = "FangSong"

# All CJK Unicode ranges that should receive the East Asian font.
_CJK_RANGES = (
    "\u3000-\u303f"   # CJK Symbols and Punctuation  (。、…《》【】)
    "\u3400-\u4dbf"   # CJK Extension A
    "\u4e00-\u9fff"   # CJK Unified Ideographs (main block)
    "\uf900-\ufaff"   # CJK Compatibility Ideographs
    "\ufe30-\ufe4f"   # CJK Compatibility Forms
    "\uff00-\uffef"   # Halfwidth and Fullwidth Forms  (，。！？：；)
    "\U00020000-\U0002a6df"  # CJK Extension B
    "\U0002a700-\U0002b73f"  # CJK Extension C
    "\U0002b740-\U0002b81f"  # CJK Extension D
    "\U0002b820-\U0002ceaf"  # CJK Extension E
    "\U0002ceb0-\U0002ebef"  # CJK Extension F
    "\U0002f800-\U0002fa1f"  # CJK Compatibility Supplement
)

CJK_PATTERN: re.Pattern[str] = re.compile(f"[{_CJK_RANGES}]")


def has_chinese(text: str) -> bool:
    """Return True if *text* contains at least one CJK character.

    Args:
        text: The string to test.

    Returns:
        True when a CJK character is found, False otherwise.
    """
    return bool(CJK_PATTERN.search(text))


def set_eastasia_font(run: object, font_name: str) -> None:
    """Set the `w:eastAsia` attribute on a run's `w:rFonts` element.

    This targets only East Asian (CJK) characters within the run; Latin
    and other character categories keep their original fonts.

    Args:
        run: A `docx.text.run.Run` instance.
        font_name: The East Asian font name to apply (e.g. ``"FangSong"``).
    """
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)


def process_paragraphs(paragraphs: list, font_name: str) -> int:
    """Apply *font_name* to every run containing CJK characters.

    Args:
        paragraphs: An iterable of `docx.text.paragraph.Paragraph` objects.
        font_name: The East Asian font name to apply.

    Returns:
        The number of runs that were modified.
    """
    count = 0
    for paragraph in paragraphs:
        for run in paragraph.runs:
            if run.text and has_chinese(run.text):
                set_eastasia_font(run, font_name)
                count += 1
    return count


def process_table(table: Table, font_name: str) -> int:
    """Recursively process all paragraphs in *table*, including nested tables.

    Args:
        table: A `docx.table.Table` instance.
        font_name: The East Asian font name to apply.

    Returns:
        The number of runs that were modified.
    """
    count = 0
    for row in table.rows:
        for cell in row.cells:
            count += process_paragraphs(cell.paragraphs, font_name)
            for nested_table in cell.tables:
                count += process_table(nested_table, font_name)
    return count


def process_document(doc: Document, font_name: str) -> int:
    """Orchestrate font changes across all document locations.

    Covers body paragraphs, all tables (nested included), headers,
    footers, and text boxes (``w:txbxContent`` elements).

    Args:
        doc: An open `docx.Document` instance.
        font_name: The East Asian font name to apply.

    Returns:
        The total number of runs that were modified.
    """
    total = 0

    # Body paragraphs
    total += process_paragraphs(doc.paragraphs, font_name)

    # Body tables (recurse for nesting)
    for table in doc.tables:
        total += process_table(table, font_name)

    # Headers and footers across all sections
    for section in doc.sections:
        for hf in (
            section.header,
            section.first_page_header,
            section.even_page_header,
            section.footer,
            section.first_page_footer,
            section.even_page_footer,
        ):
            if hf is not None:
                total += process_paragraphs(hf.paragraphs, font_name)
                for table in hf.tables:
                    total += process_table(table, font_name)

    # Text boxes — search the body XML for w:txbxContent elements
    body = doc.element.body
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    for txbx in body.findall(".//w:txbxContent", ns):
        for para_elem in txbx.findall(".//w:p", ns):
            # Wrap raw XML paragraphs using python-docx Paragraph proxy
            from docx.text.paragraph import Paragraph

            para = Paragraph(para_elem, doc)
            total += process_paragraphs([para], font_name)

    return total


def main() -> None:
    """CLI entry point: parse arguments and run the font changer."""
    parser = argparse.ArgumentParser(
        description="Apply an East Asian font to all Chinese characters in a .docx file."
    )
    parser.add_argument("input", help="Path to the source .docx file.")
    parser.add_argument(
        "--output",
        default=None,
        help=(
            "Path for the modified .docx file. "
            "Defaults to <input_stem>_modified.docx in the same directory."
        ),
    )
    parser.add_argument(
        "--font",
        default=DEFAULT_FONT,
        help=f"East Asian font name to apply (default: {DEFAULT_FONT}).",
    )
    args = parser.parse_args()

    input_path = Path(args.input)
    if not input_path.exists():
        print(f"Error: file not found: {input_path}", file=sys.stderr)
        sys.exit(1)
    if input_path.suffix.lower() != ".docx":
        print(f"Error: input must be a .docx file, got: {input_path}", file=sys.stderr)
        sys.exit(1)

    output_path = (
        Path(args.output)
        if args.output
        else input_path.with_name(f"{input_path.stem}_modified.docx")
    )

    print(f"Opening:  {input_path}")
    doc = Document(str(input_path))

    modified = process_document(doc, args.font)

    doc.save(str(output_path))
    print(f"Saved:    {output_path}")
    print(f"Runs modified: {modified}  (font: {args.font})")


if __name__ == "__main__":
    main()
